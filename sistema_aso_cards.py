import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os

# =========================================================
# CONFIGURAÇÃO DA PÁGINA
# =========================================================

st.set_page_config(
    page_title="Gestão ASO Pro - Junior 360",
    layout="wide",
    page_icon="🛡️"
)

# =========================================================
# CONFIGURAÇÕES
# =========================================================

SHEET_ID = "1G_oVT9gK-n_jGh5R4g65qUwK_MfQGvCX-SA4NHNNflU"

UNIDADES = {
    "D-ITU": "1323067532",
    "D-MG": "1071212860",
    "J-CHP": "1549718037",
    "S-SJ": "1712391604",
    "J-CTBA": "145843404"
}

# LOGO SIDEBAR
PATH_LOGO = r"C:\Users\dilceu.gomes\Desktop\sistema_aso\logo.png"

# LOGO DOCUMENTO WORD
PATH_LOGO_DOC = r"C:\Users\dilceu.gomes\Desktop\sistema_aso\adivitta.png"

# =========================================================
# CSS GLOBAL
# =========================================================

st.markdown("""
<style>

/* FUNDO */
.stApp {
    background-color: #f1f5f9;
}

/* SIDEBAR */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #8390a8, #1e293b);
}

/* TEXTO SIDEBAR */
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stMarkdown,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: white !important;
}

/* LABEL SELECT */
label[data-testid="stWidgetLabel"] p {
    color: white !important;
    font-weight: bold !important;
    font-size: 16px !important;
}

/* SELECTBOX */
div[data-baseweb="select"] > div {
    background-color: white !important;
    border-radius: 12px !important;
    border: 2px solid #2563eb !important;
    min-height: 50px !important;
}

/* TEXTO SELECT */
div[data-baseweb="select"] * {
    color: black !important;
}

/* TEXTO INTERNO */
div[data-baseweb="select"] span {
    color: black !important;
    font-weight: 700 !important;
    font-size: 15px !important;
}

/* INPUT */
div[data-baseweb="select"] input {
    color: black !important;
    caret-color: black !important;
}

/* DROPDOWN */
ul {
    background-color: white !important;
}

li {
    color: black !important;
}

/* MÉTRICAS */
div[data-testid="metric-container"] {
    background: white;
    border-radius: 18px;
    padding: 15px;
    border: 1px solid #e2e8f0;
    box-shadow: 0 4px 18px rgba(0,0,0,0.06);
}

/* BOTÕES */
.stDownloadButton button {
    width: 100%;
    background: linear-gradient(90deg, #2563eb, #1d4ed8);
    color: white;
    border: none;
    border-radius: 12px;
    padding: 12px;
    font-weight: bold;
}

.stDownloadButton button:hover {
    background: linear-gradient(90deg, #1d4ed8, #1e40af);
}

/* EXPANDER */
.streamlit-expanderHeader {
    font-weight: bold;
}

/* TÍTULOS */
.main-title {
    font-size: 40px;
    font-weight: 800;
    color: #111827;
}

.sub-title {
    color: #64748b;
    margin-bottom: 20px;
}

</style>
""", unsafe_allow_html=True)

# =========================================================
# FUNÇÃO LEITURA PLANILHA
# =========================================================

@st.cache_data(ttl=300)
def load_data(gid):

    url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv&gid={gid}"

    df = pd.read_csv(url)

    df.columns = df.columns.str.strip()

    return df

# =========================================================
# GERAR DOCUMENTO WORD
# =========================================================

def gerar_docx(dados, tipo, data_sugestao):

    doc = Document()

    # LOGO DOCUMENTO
    if os.path.exists(PATH_LOGO_DOC):

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.add_run().add_picture(
            PATH_LOGO_DOC,
            width=Inches(1.2)
        )

    # TÍTULO
    titulo = doc.add_paragraph()

    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = titulo.add_run(
        f"FORMULÁRIO PARA AGENDAMENTO {tipo}"
    )

    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("")

    # UNIDADE DA PLANILHA
    unidade_colaborador = str(dados["UNIDADE"])

    # TABELA
    table = doc.add_table(rows=7, cols=2)

    table.style = "Table Grid"

    labels = [
        "Nome Completo",
        "Cargo",
        "Setor",
        "Unidade",
        "Cliente",
        "Local",
        "Data Sugestão"
    ]

    valores = [
        str(dados["Nome"]),
        str(dados["Cargo"]),
        str(dados["Setor"]),
        unidade_colaborador,
        "MACROMAQ",
        "Arapoti",
        data_sugestao.strftime("%d/%m/%Y")
    ]

    for i, (l, v) in enumerate(zip(labels, valores)):

        table.rows[i].cells[0].text = l
        table.rows[i].cells[1].text = v

        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    target = BytesIO()

    doc.save(target)

    return target.getvalue()

# =========================================================
# TOPO
# =========================================================

st.markdown("""
<div class="main-title">
🛡️ Gestão ASO Pro
</div>

<div class="sub-title">
Sistema Inteligente de Gestão Ocupacional •
</div>
""", unsafe_allow_html=True)

# =========================================================
# SIDEBAR
# =========================================================

if os.path.exists(PATH_LOGO):
    st.sidebar.image(PATH_LOGO, width=300)

aba_nome = st.sidebar.selectbox(
    "🏢 Selecione a unidade",
    list(UNIDADES.keys())
)

st.sidebar.success("✅ Sistema Online")

# =========================================================
# CARREGAR DADOS
# =========================================================

try:

    df = load_data(UNIDADES[aba_nome])

except Exception as e:

    st.error(f"Erro ao carregar dados:\n\n{e}")

    st.stop()

# =========================================================
# PROCESSAMENTO
# =========================================================

if df.empty:

    st.warning("Nenhum dado encontrado.")

else:

    hoje = datetime.now()

    prazo_limite = hoje + timedelta(days=10)

    # DATAS
    df["Venc"] = pd.to_datetime(
        df["Venc"],
        dayfirst=True,
        errors="coerce"
    )

    df = df.dropna(subset=["Venc"])

    # DIAS
    df["Dias"] = (df["Venc"] - hoje).dt.days

    # ALERTAS
    alertas = df[
        df["Venc"] <= prazo_limite
    ].copy().sort_values(by="Venc")

    total_vencidos = len(alertas[alertas["Dias"] < 0])

    # KPIs
    c1, c2, c3, c4 = st.columns(4)

    c1.metric("🏢 Unidade", aba_nome)
    c2.metric("👥 Colaboradores", len(df))
    c3.metric("⚠️ Pendentes", len(alertas))
    c4.metric("🚨 Vencidos", total_vencidos)

    st.markdown("---")

    st.subheader("📋 Colaboradores que precisam de atenção")

    # CARDS
    cols = st.columns(2)

    for idx, (_, row) in enumerate(alertas.iterrows()):

        col = cols[idx % 2]

        dias = int(row["Dias"])

        # STATUS
        if dias < 0:

            cor_borda = "#ef4444"
            fundo = "#fff1f2"
            status = "🚨 ASO VENCIDO"
            cor_status = "#ef4444"

        elif dias <= 3:

            cor_borda = "#f59e0b"
            fundo = "#fff7ed"
            status = f"⚠️ Vence em {dias} dias"
            cor_status = "#f59e0b"

        else:

            cor_borda = "#10b981"
            fundo = "#ecfdf5"
            status = f"✅ Vence em {dias} dias"
            cor_status = "#10b981"

        # CARD HTML
        html_card = f"""
        <div style="
            background:{fundo};
            border-left:8px solid {cor_borda};
            border-radius:18px;
            padding:22px;
            margin-bottom:15px;
            box-shadow:0 4px 18px rgba(0,0,0,0.08);
            font-family:Arial;
        ">

            <div style="
                font-size:22px;
                font-weight:700;
                color:#111827;
                margin-bottom:10px;
            ">
                {row['Nome']}
            </div>

            <div style="
                color:#475569;
                font-size:15px;
                line-height:1.8;
            ">
                👔 <b>Cargo:</b> {row['Cargo']}<br>
                🏭 <b>Setor:</b> {row['Setor']}<br>
                🏢 <b>Unidade:</b> {row['UNIDADE']}<br>
                📅 <b>Vencimento:</b> {row['Venc'].strftime('%d/%m/%Y')}
            </div>

            <div style="
                margin-top:15px;
                font-size:16px;
                font-weight:bold;
                color:{cor_status};
            ">
                {status}
            </div>

        </div>
        """

        with col:

            components.html(
                html_card,
                height=250
            )

            with st.expander(
                f"📄 Gerar Agendamento - {row['Nome'].split()[0]}"
            ):

                tipo = st.selectbox(
                    "Tipo de Exame",
                    [
                        "PERIÓDICO",
                        "MUDANÇA DE RISCO",
                        "RETORNO"
                    ],
                    key=f"tipo_{idx}"
                )

                dt_sug = st.date_input(
                    "Data sugerida",
                    value=hoje + timedelta(days=2),
                    key=f"data_{idx}"
                )

                btn_doc = gerar_docx(
                    row,
                    tipo,
                    dt_sug
                )

                st.download_button(
                    label="📥 Baixar Documento",
                    data=btn_doc,
                    file_name=f"ASO_{row['Nome']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"btn_{idx}"
                )